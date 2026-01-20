"""
Excel ve DOCX dosyalarından verileri okuyup veritabanına entegre eder
"""
import pandas as pd
from docx import Document
import sqlite3
import datetime
import os
import sys

# Veritabanı yolu - EXE ve Python için optimize
if getattr(sys, 'frozen', False):
    # EXE modunda: EXE'nin bulunduğu klasör
    APP_PATH = os.path.dirname(sys.executable)
else:
    # Python modunda: Script'in bulunduğu klasör
    APP_PATH = os.path.dirname(os.path.abspath(__file__))

DB_NAME = os.path.join(APP_PATH, "leta_data.db")

# Veri dosyalarını bulmak için arama yolları
def veri_dosyasi_bul(dosya_adi):
    """Veri dosyasını farklı konumlarda ara"""
    arama_yollari = [
        dosya_adi,  # Mevcut klasör
        os.path.join(APP_PATH, dosya_adi),  # APP_PATH'te
        os.path.join(APP_PATH, "veriler", dosya_adi),  # veriler klasöründe
        os.path.join(APP_PATH, "..", "veriler", dosya_adi),  # Bir üst klasörde veriler
        os.path.join(os.path.dirname(APP_PATH), "veriler", dosya_adi),  # Üst klasörde veriler
    ]
    
    for yol in arama_yollari:
        if os.path.exists(yol):
            return yol
    return None

def veritabani_baglan():
    return sqlite3.connect(DB_NAME)

def docx_ogrenci_listesi_oku():
    """ÖĞRENCİ LİSTESİ.docx dosyasından öğrenci adlarını oku"""
    try:
        dosya_yolu = veri_dosyasi_bul("ÖĞRENCİ LİSTESİ.docx")
        if not dosya_yolu:
            print("⚠ ÖĞRENCİ LİSTESİ.docx dosyası bulunamadı")
            return []
        doc = Document(dosya_yolu)
        ogrenciler = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) > 2 and not text.startswith("ÖĞRENCİ"):
                # Boş satırları ve başlıkları filtrele
                if text and text.isupper() or any(c.isalpha() for c in text):
                    ogrenciler.append(text.upper())
        return ogrenciler
    except Exception as e:
        print(f"DOCX okuma hatası: {e}")
        return []

def docx_aile_numaralari_oku():
    """ÖĞRENCİ AİLE NUMARALARI.docx dosyasından öğrenci ve aile bilgilerini oku"""
    try:
        dosya_yolu = veri_dosyasi_bul("ÖĞRENCİ AİLE NUMARALARI.docx")
        if not dosya_yolu:
            print("⚠ ÖĞRENCİ AİLE NUMARALARI.docx dosyası bulunamadı")
            return {}
        doc = Document(dosya_yolu)
        ogrenci_bilgileri = {}
        current_ogrenci = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Öğrenci adı (büyük harfle başlayan, iki nokta üst üste ile biten)
            if ":" in text and len(text.split(":")[0]) > 3:
                current_ogrenci = text.split(":")[0].strip().upper()
                if current_ogrenci not in ogrenci_bilgileri:
                    ogrenci_bilgileri[current_ogrenci] = {
                        'veli_adi': '',
                        'veli_telefon': '',
                        'telefon': '',
                        'notlar': ''
                    }
            elif current_ogrenci:
                # Telefon numarası kontrolü
                if any(char.isdigit() for char in text) and len(text) > 8:
                    # Anne, Baba, vb. kontrolü
                    if 'anne' in text.lower() or 'baba' in text.lower() or 'veli' in text.lower():
                        parts = text.split(":")
                        if len(parts) >= 2:
                            veli_adi = parts[0].strip()
                            telefon = parts[1].strip().replace(" ", "").replace("-", "")
                            if not ogrenci_bilgileri[current_ogrenci]['veli_adi']:
                                ogrenci_bilgileri[current_ogrenci]['veli_adi'] = veli_adi
                                ogrenci_bilgileri[current_ogrenci]['veli_telefon'] = telefon
                    else:
                        # Sadece telefon numarası
                        telefon = text.replace(" ", "").replace("-", "").replace("(", "").replace(")", "")
                        if len(telefon) >= 10:
                            if not ogrenci_bilgileri[current_ogrenci]['telefon']:
                                ogrenci_bilgileri[current_ogrenci]['telefon'] = telefon
        
        return ogrenci_bilgileri
    except Exception as e:
        print(f"DOCX okuma hatası: {e}")
        return {}

def excel_seans_okuma():
    """SEANS ÜCRET TAKİP.xlsx dosyasından seans ve ücret bilgilerini oku"""
    try:
        dosya_yolu = veri_dosyasi_bul("SEANS ÜCRET TAKİP.xlsx")
        if not dosya_yolu:
            print("⚠ SEANS ÜCRET TAKİP.xlsx dosyası bulunamadı")
            return []
        df = pd.read_excel(dosya_yolu)
        seanslar = []
        
        for idx, row in df.iterrows():
            if pd.notna(row.get('TARİH')) and pd.notna(row.get('DANIŞANIN ADI')):
                seans = {
                    'tarih': row.get('TARİH'),
                    'danisan_adi': str(row.get('DANIŞANIN ADI', '')).strip().upper(),
                    'terapist': str(row.get('TERAPİSTİN ADI', '')).strip(),
                    'alinan_ucret': row.get('ALINAN ÜCRET', 0),
                    'alinacak_ucret': row.get('ALINACAK ÜCRET', 0),
                    'odeme_sekli': str(row.get('ÖDEME ŞEKLİ', '')),
                    'guncel': str(row.get('GÜNCEL', ''))
                }
                # Tarih formatını düzelt
                if isinstance(seans['tarih'], pd.Timestamp):
                    seans['tarih'] = seans['tarih'].strftime('%Y-%m-%d')
                elif isinstance(seans['tarih'], str):
                    try:
                        dt = pd.to_datetime(seans['tarih'])
                        seans['tarih'] = dt.strftime('%Y-%m-%d')
                    except:
                        pass
                
                # Ücretleri sayıya çevir
                try:
                    if pd.notna(seans['alinan_ucret']):
                        seans['alinan_ucret'] = float(str(seans['alinan_ucret']).replace(',', '.'))
                    else:
                        seans['alinan_ucret'] = 0
                except:
                    seans['alinan_ucret'] = 0
                
                try:
                    if pd.notna(seans['alinacak_ucret']):
                        seans['alinacak_ucret'] = float(str(seans['alinacak_ucret']).replace(',', '.'))
                    else:
                        seans['alinacak_ucret'] = 0
                except:
                    seans['alinacak_ucret'] = 0
                
                seanslar.append(seans)
        
        return seanslar
    except Exception as e:
        print(f"Excel okuma hatası: {e}")
        return []

def veritabani_tablolari_olustur():
    """Veritabanı tablolarını oluştur"""
    conn = veritabani_baglan()
    cursor = conn.cursor()
    
    # Danışanlar tablosu
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS danisanlar (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ad_soyad TEXT NOT NULL,
            telefon TEXT,
            email TEXT,
            adres TEXT,
            dogum_tarihi TEXT,
            veli_adi TEXT,
            veli_telefon TEXT,
            notlar TEXT,
            olusturma_tarihi TEXT,
            aktif INTEGER DEFAULT 1
        )
    """)
    
    # Seans kayıtları tablosu (mevcut)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS kayitlar (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            tarih TEXT,
            danisan_adi TEXT,
            terapist TEXT,
            hizmet_bedeli REAL,
            alinan_ucret REAL,
            kalan_borc REAL,
            notlar TEXT,
            son_islem_tarihi TEXT
        )
    """)
    
    conn.commit()
    conn.close()
    print("Veritabanı tabloları oluşturuldu/kontrol edildi")

def veritabanina_entegre_et():
    """Tüm verileri veritabanına entegre et"""
    # Önce tabloları oluştur
    veritabani_tablolari_olustur()
    
    conn = veritabani_baglan()
    cursor = conn.cursor()
    
    print("\n" + "="*80)
    print("VERİ ENTEGRASYONU BAŞLIYOR...")
    print("="*80)
    
    # 1. Öğrenci listesini danışanlar tablosuna ekle
    print("\n1. Öğrenci listesi işleniyor...")
    ogrenciler = docx_ogrenci_listesi_oku()
    aile_bilgileri = docx_aile_numaralari_oku()
    
    danisan_sayisi = 0
    for ogrenci_adi in ogrenciler:
        # Zaten var mı kontrol et
        cursor.execute("SELECT COUNT(*) FROM danisanlar WHERE ad_soyad = ?", (ogrenci_adi,))
        if cursor.fetchone()[0] == 0:
            bilgi = aile_bilgileri.get(ogrenci_adi, {})
            cursor.execute("""
                INSERT INTO danisanlar (ad_soyad, telefon, veli_adi, veli_telefon, notlar, olusturma_tarihi, aktif)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
                ogrenci_adi,
                bilgi.get('telefon', ''),
                bilgi.get('veli_adi', ''),
                bilgi.get('veli_telefon', ''),
                bilgi.get('notlar', ''),
                datetime.datetime.now().strftime('%Y-%m-%d %H:%M'),
                1
            ))
            danisan_sayisi += 1
    
    print(f"   {danisan_sayisi} yeni danışan eklendi")
    
    # 2. Excel'den seans verilerini ekle
    print("\n2. Seans verileri işleniyor...")
    seanslar = excel_seans_okuma()
    
    seans_sayisi = 0
    for seans in seanslar:
        # Danışan ID'sini bul
        cursor.execute("SELECT id FROM danisanlar WHERE ad_soyad = ?", (seans['danisan_adi'],))
        danisan_result = cursor.fetchone()
        
        if danisan_result:
            danisan_id = danisan_result[0]
        else:
            # Danışan yoksa ekle
            cursor.execute("""
                INSERT INTO danisanlar (ad_soyad, olusturma_tarihi, aktif)
                VALUES (?, ?, ?)
            """, (seans['danisan_adi'], datetime.datetime.now().strftime('%Y-%m-%d %H:%M'), 1))
            danisan_id = cursor.lastrowid
        
        # Seans kaydı var mı kontrol et (tarih + danışan + terapist)
        cursor.execute("""
            SELECT COUNT(*) FROM kayitlar 
            WHERE tarih = ? AND danisan_adi = ? AND terapist = ?
        """, (seans['tarih'], seans['danisan_adi'], seans['terapist']))
        
        if cursor.fetchone()[0] == 0:
            # Ücret hesaplama
            hizmet_bedeli = seans['alinacak_ucret'] if seans['alinacak_ucret'] > 0 else seans['alinan_ucret']
            alinan_ucret = seans['alinan_ucret'] if seans['alinan_ucret'] > 0 else 0
            kalan_borc = hizmet_bedeli - alinan_ucret
            
            cursor.execute("""
                INSERT INTO kayitlar (tarih, danisan_adi, terapist, hizmet_bedeli, alinan_ucret, kalan_borc, notlar, son_islem_tarihi)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                seans['tarih'],
                seans['danisan_adi'],
                seans['terapist'],
                hizmet_bedeli,
                alinan_ucret,
                kalan_borc,
                f"Ödeme: {seans['odeme_sekli']}, Güncel: {seans['guncel']}",
                datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
            ))
            seans_sayisi += 1
    
    print(f"   {seans_sayisi} yeni seans kaydı eklendi")
    
    conn.commit()
    conn.close()
    
    print("\n" + "="*80)
    print("VERİ ENTEGRASYONU TAMAMLANDI!")
    print("="*80)
    print(f"\nÖzet:")
    print(f"- {danisan_sayisi} yeni danışan eklendi")
    print(f"- {seans_sayisi} yeni seans kaydı eklendi")
    print(f"\nSistem yeni veri girişi için hazır!")

if __name__ == "__main__":
    veritabanina_entegre_et()

