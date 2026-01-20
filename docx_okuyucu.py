"""
DOCX dosyalarını okumak için hazır script
Kullanıcı DOCX dosyalarını eklediğinde bu script ile analiz edilecek
"""
try:
    from docx import Document
    import sys
    import os
    
    # Klasördeki tüm DOCX dosyalarını bul
    docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]
    
    if not docx_files:
        print("DOCX dosyası bulunamadı!")
        sys.exit(1)
    
    for dosya in docx_files:
        print(f"\n{'='*80}")
        print(f"DOSYA: {dosya}")
        print('='*80)
        
        try:
            doc = Document(dosya)
            
            # Tüm paragrafları oku
            print("\nPARAGRAFLAR:")
            for i, para in enumerate(doc.paragraphs[:50], 1):  # İlk 50 paragraf
                if para.text.strip():
                    print(f"{i}. {para.text[:200]}")  # İlk 200 karakter
            
            # Tabloları oku
            print(f"\nTABLO SAYISI: {len(doc.tables)}")
            for i, table in enumerate(doc.tables, 1):
                print(f"\n--- TABLO {i} ---")
                for row_idx, row in enumerate(table.rows[:10], 1):  # İlk 10 satır
                    row_data = [cell.text.strip() for cell in row.cells]
                    print(f"Satır {row_idx}: {' | '.join(row_data)}")
            
        except Exception as e:
            print(f"HATA: {e}")
    
    print("\n" + "="*80)
    print("ANALİZ TAMAMLANDI")
    print("="*80)
    
except ImportError:
    print("python-docx kütüphanesi yüklü değil!")
    print("Yüklemek için: pip install python-docx")
    import sys
    sys.exit(1)

